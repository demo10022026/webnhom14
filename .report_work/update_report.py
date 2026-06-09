from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "bao_cao_shopvn.docx"
BACKUP = ROOT / ".report_work" / "bao_cao_shopvn_before_update.docx"
ASSETS = ROOT / ".report_work" / "assets"
SHOTS = ROOT / ".report_work" / "screenshots"


def set_text(paragraph, text):
    paragraph.clear()
    paragraph.add_run(text)
    return paragraph


def insert_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = paragraph._parent.add_paragraph()
    p._p.getparent().remove(p._p)
    new_p.getparent().replace(new_p, p._p)
    if style:
        p.style = style
    if text:
        p.add_run(text)
    return p


def find_para(doc, predicate, label):
    for p in doc.paragraphs:
        if predicate(p.text.strip()):
            return p
    raise ValueError(f"Not found: {label}")


def find_starts(doc, prefix):
    return find_para(doc, lambda text: text.startswith(prefix), prefix)


def find_contains(doc, needle):
    return find_para(doc, lambda text: needle in text, needle)


def find_after(doc, anchor_prefix, target_prefix):
    found_anchor = False
    for p in doc.paragraphs:
        text = p.text.strip()
        if not found_anchor and text.startswith(anchor_prefix):
            found_anchor = True
            continue
        if found_anchor and text.startswith(target_prefix):
            return p
    raise ValueError(f"Not found after {anchor_prefix}: {target_prefix}")


def paragraph_index(doc, paragraph):
    target = paragraph._p
    for index, candidate in enumerate(doc.paragraphs):
        if candidate._p is target:
            return index
    raise ValueError("Paragraph is not in document paragraph list")


def clear_picture_paragraph(paragraph, image_path, width=6.2):
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width))


def main():
    shutil.copy2(REPORT, BACKUP)
    doc = Document(REPORT)

    # Replace the four old database diagram images with the current ERD slices.
    for idx in range(1, 5):
        caption = find_after(doc, "CHƯƠNG 4:", f"Hình 4.{idx}.")
        caption_index = paragraph_index(doc, caption)
        image_paragraph = doc.paragraphs[caption_index - 1]
        clear_picture_paragraph(image_paragraph, ASSETS / f"defaultdb_part_{idx}.png", width=6.2)
        set_text(caption, f"Hình 4.{idx}. Sơ đồ cơ sở dữ liệu ShopVN cập nhật - phần {idx}")
        caption.style = "CaptionVN"
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    set_text(
        find_starts(doc, "Cơ sở dữ liệu ecommerce_db được thiết kế"),
        "Cơ sở dữ liệu ecommerce_db được thiết kế cho mô hình thương mại điện tử nhiều người bán. "
        "Schema hiện tại sử dụng quy ước snake_case và chia thành các nhóm bảng chính: tài khoản - phân quyền, "
        "seller/shop, sản phẩm - biến thể - ảnh, giỏ hàng - đơn hàng - thanh toán, địa chỉ, voucher/flash sale, "
        "đánh giá, chat, thông báo, OTP và refresh token. Bản cập nhật đã bổ sung rõ các bảng phục vụ seller onboarding, "
        "thông tin ngân hàng người bán, giao dịch thanh toán và đối soát đơn hàng.",
    )

    p = find_starts(doc, "Cần tránh lưu các giá trị bí mật")
    p = insert_after(
        p,
        "Bảng payments hiện lưu phương thức thanh toán, mã giao dịch, trạng thái xác nhận và thời điểm tạo giao dịch; "
        "kết hợp với orders để xử lý COD và thanh toán SePay/VietQR.",
        "List Bullet",
    )
    insert_after(
        p,
        "Các bảng user_addresses lưu đầy đủ mã/tên tỉnh, quận/huyện, phường/xã theo API địa chỉ phân cấp, "
        "giúp checkout hiển thị lại địa chỉ chính xác sau khi tạo mới.",
        "List Bullet",
    )

    # Add current screenshots captured from the local frontend.
    caption_512 = find_after(doc, "CHƯƠNG 5:", "Hình 5.12.")
    img = insert_after(caption_512, "", "NoIndent")
    clear_picture_paragraph(img, SHOTS / "01-home.png", width=6.2)
    cap = insert_after(img, "Hình 5.13. Trang chủ ShopVN phiên bản hiện tại chụp từ môi trường local", "CaptionVN")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img = insert_after(cap, "", "NoIndent")
    clear_picture_paragraph(img, SHOTS / "02-login.png", width=5.4)
    cap = insert_after(img, "Hình 5.14. Giao diện đăng nhập phiên bản hiện tại", "CaptionVN")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img = insert_after(cap, "", "NoIndent")
    clear_picture_paragraph(img, SHOTS / "03-register.png", width=5.4)
    cap = insert_after(img, "Hình 5.15. Giao diện đăng ký tài khoản phiên bản hiện tại", "CaptionVN")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    set_text(
        find_starts(doc, "Giỏ hàng hiển thị các sản phẩm đã chọn"),
        "Giỏ hàng hiển thị sản phẩm đã chọn, biến thể, đơn giá, số lượng và tổng tiền. Khi đặt hàng, hệ thống xử lý "
        "địa chỉ nhận hàng, phí vận chuyển, voucher và phương thức thanh toán. Với thanh toán SePay/VietQR, frontend "
        "tạo giao dịch thanh toán, hiển thị mã QR theo thông tin đơn hàng và polling trạng thái thanh toán để chuyển "
        "người dùng về trang Đơn mua của tôi sau khi webhook xác nhận đã thanh toán.",
    )
    set_text(
        find_starts(doc, "Các trang hồ sơ cá nhân"),
        "Các trang hồ sơ cá nhân, địa chỉ, voucher và lịch sử đơn hàng giúp người dùng quản lý thông tin cá nhân và "
        "theo dõi giao dịch đã phát sinh. Bản mới đã bổ sung upload ảnh đại diện người dùng qua Cloudinary, modal thêm "
        "địa chỉ theo luồng tỉnh - quận/huyện - phường/xã, và hiển thị lại địa chỉ ngay trong trang thanh toán sau khi tạo mới.",
    )
    set_text(
        find_starts(doc, "Giao diện seller gồm đăng ký"),
        "Giao diện seller gồm đăng ký trở thành người bán, theo dõi trạng thái hồ sơ, upload giấy tờ, thiết lập shop "
        "và cấu hình thông tin nhận tiền. Sau khi được duyệt, seller có thể quản lý sản phẩm, tồn kho, đơn hàng và trao đổi "
        "với khách hàng. Phần hội thoại phía seller đã được điều chỉnh để hiển thị tên khách hàng thay vì tên shop của chính seller.",
    )
    set_text(
        find_starts(doc, "Giao diện quản trị dành cho admin/manager"),
        "Giao diện quản trị dành cho admin/manager. Các chức năng chính gồm dashboard thống kê, quản lý sản phẩm, xem chi tiết "
        "sản phẩm, cập nhật trạng thái sản phẩm, duyệt seller, quản lý người dùng theo quyền và theo dõi dữ liệu kinh doanh. "
        "Hệ thống phân quyền để manager chỉ thao tác trong phạm vi được phép, còn admin có quyền quản trị cao nhất.",
    )

    # Chapter 6 additions.
    old_expansion = find_starts(doc, "6.7 Định hướng mở rộng")
    set_text(old_expansion, "6.9 Định hướng mở rộng")
    old_expansion.style = "Heading 2"

    dashboard_text = find_starts(doc, "Dashboard admin tổng hợp dữ liệu")
    h_payment = insert_after(dashboard_text, "6.7 Thanh toán SePay/VietQR và xác nhận realtime", "Heading 2")
    payment_text = insert_after(
        h_payment,
        "Project mới đã chuyển từ mô phỏng thủ công sang luồng thanh toán online qua SePay/VietQR. Khi người dùng chọn "
        "thanh toán online, backend tạo payment gắn với order, frontend hiển thị QR chứa số tiền và nội dung chuyển khoản, "
        "sau đó polling endpoint trạng thái. SePay gửi webhook về backend; backend xác thực API key, đối chiếu mã đơn/nội dung "
        "giao dịch, cập nhật payment thành PAID và chuyển đơn hàng sang trạng thái đã thanh toán. Nút mô phỏng thanh toán đã được "
        "loại khỏi giao diện người dùng để luồng demo gần với thực tế hơn.",
        "Normal",
    )
    h_deploy = insert_after(payment_text, "6.8 Triển khai backend và webhook", "Heading 2")
    insert_after(
        h_deploy,
        "Backend đã được cấu hình để triển khai trên Railway với public URL phục vụ frontend và webhook SePay. Trong giai đoạn "
        "phát triển local, Cloudflare Tunnel được dùng để tạo URL HTTPS tạm thời trỏ về backend localhost:8080, giúp SePay test mode "
        "gửi webhook vào máy dev. Các cấu hình nhạy cảm như database, mail, JWT, Cloudinary và SePay API key được đưa vào biến môi trường "
        "thay vì hard-code trong mã nguồn.",
        "Normal",
    )

    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("Hoàn thiện quy trình checkout và thanh toán trực tuyến"):
            set_text(
                paragraph,
                "Mở rộng thanh toán sang đối soát tự động, hoàn tiền, lưu lịch sử webhook và hỗ trợ thêm cổng thanh toán khác "
                "khi đề tài cần so sánh nhiều phương thức.",
            )
            paragraph.style = "List Bullet"
            break

    # Chapter 7 additions.
    old_result = find_starts(doc, "7.9 Đánh giá kết quả kiểm thử")
    set_text(old_result, "7.10 Đánh giá kết quả kiểm thử")
    old_result.style = "Heading 2"
    insert_anchor = doc.paragraphs[paragraph_index(doc, old_result) - 1]
    h_test = insert_after(insert_anchor, "7.9 Kiểm thử thanh toán SePay và webhook", "Heading 2")
    insert_after(
        h_test,
        "Các kịch bản kiểm thử bổ sung gồm: tạo đơn hàng với phương thức SePay, hiển thị mã QR đúng số tiền và nội dung chuyển khoản, "
        "nhận webhook test mode từ SePay qua Cloudflare Tunnel/Railway, cập nhật trạng thái payment/order sau khi giao dịch hợp lệ, "
        "polling frontend phát hiện trạng thái đã thanh toán và tự chuyển về trang Đơn mua của tôi. Trường hợp sai API key, sai mã đơn "
        "hoặc giao dịch không khớp số tiền cần bị từ chối và không cập nhật đơn hàng.",
        "Normal",
    )

    set_text(
        find_starts(doc, "Qua các trường hợp kiểm thử chức năng"),
        "Qua các trường hợp kiểm thử chức năng, hệ thống đáp ứng được các nghiệp vụ nền tảng của website thương mại điện tử. "
        "Các luồng đăng ký, đăng nhập, phân quyền, địa chỉ, giỏ hàng, đặt hàng, seller, admin/manager và thanh toán SePay ở môi trường "
        "demo đã được kiểm tra theo các kịch bản chính. Khi triển khai production cần tiếp tục bổ sung kiểm thử tải, kiểm thử bảo mật webhook, "
        "kiểm thử giao dịch trùng và giám sát lỗi runtime.",
    )

    set_text(
        find_starts(doc, "Hạn chế hiện tại là một số chức năng"),
        "Hạn chế hiện tại là hệ thống vẫn cần hoàn thiện sâu hơn ở các phần vận chuyển thực tế, hoàn tiền/đối soát nâng cao, "
        "gợi ý sản phẩm, tối ưu tìm kiếm khi dữ liệu lớn và kiểm thử tải. Phần thanh toán online đã có luồng SePay/VietQR ở mức "
        "phù hợp cho đồ án và môi trường demo, nhưng khi đưa vào vận hành thật cần bổ sung logging webhook, chữ ký xác thực mạnh hơn, "
        "quy trình xử lý giao dịch trùng và giám sát lỗi production.",
    )
    set_text(
        find_starts(doc, "Trong tương lai, nhóm có thể mở rộng"),
        "Trong tương lai, nhóm có thể mở rộng hệ thống theo hướng chat real-time, thông báo real-time, gợi ý sản phẩm, tích hợp vận chuyển, "
        "hoàn tiền, báo cáo doanh thu nâng cao và triển khai frontend/backend trên cloud với domain cố định để người dùng có thể truy cập qua Internet.",
    )

    last = doc.paragraphs[-1]
    for text in [
        "Tài liệu SePay Webhook, VietQR và môi trường test mode.",
        "Tài liệu Railway Deployment và cấu hình biến môi trường.",
        "Tài liệu Cloudflare Tunnel dùng cho webhook local trong quá trình phát triển.",
    ]:
        last = insert_after(last, text, "NoIndent")

    doc.save(REPORT)
    print(f"saved={REPORT}")
    print(f"backup={BACKUP}")


if __name__ == "__main__":
    main()
